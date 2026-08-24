import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders {} >
            <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    '''.format(nsdecls('w')))
    tcPr.append(tcBorders)

def set_table_header_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders {} >
            <w:top w:val="single" w:sz="8" w:space="0" w:color="333333"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="333333"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    '''.format(nsdecls('w')))
    tcPr.append(tcBorders)

def create_report():
    doc = docx.Document()

    # Set page margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style definitions
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Performance Evaluation and Architecture of a Distributed WebSocket Load Balancer for Secure Real-Time Messaging")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.font.bold = True

    # Student metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(18)
    
    r_name = p_meta.add_run("Lekkala Sashank\n")
    r_name.font.bold = True
    r_name.font.size = Pt(12)
    
    r_roll = p_meta.add_run("Roll Number: 12341330 | Course: Computer Systems & Design (CSD)")
    r_roll.font.size = Pt(11)

    # Helper for headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.font.bold = True
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11.5)
        r.font.bold = True
        return p

    def add_para(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        return p

    # --- Section 1: Introduction ---
    add_heading_1("1. Introduction and Objectives")
    add_para(
        "This report presents the design, implementation, and empirical evaluation of a high-concurrency "
        "Load Balancer developed in Go. The system distributes real-time, encrypted group messaging traffic "
        "across multiple backend server instances deployed in isolated container environments."
    )
    add_para(
        "The primary objectives of this project include:\n"
        "1. Developing a custom Load Balancer in Go capable of proxying standard HTTP requests and handling full-duplex WebSocket connections.\n"
        "2. Integrating the previously developed secure messaging backend across three independent backend container instances.\n"
        "3. Building an automated WebSocket Load Generator in Go to benchmark system behavior under high concurrency (50 simultaneous virtual users).\n"
        "4. Measuring and comparing throughput, connection reliability, and latency distribution between a Single Backend deployment (Sys2) and a Three-Backend load-balanced cluster (Sys2 + Sys3 + Sys4)."
    )

    # --- Section 2: Infrastructure & Network Topology ---
    add_heading_1("2. Infrastructure and Network Topology")
    add_para(
        "The application was deployed across four allocated Docker containers on the host server 10.1.75.51. "
        "Each container operates under an isolated student user environment with a private filesystem, "
        "dedicated port mappings, and distinct internal IP addresses."
    )

    # Table 1: Systems Mapping
    table1 = doc.add_table(rows=5, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False

    headers1 = ["System", "Assigned Role", "SSH Command", "External Host", "Container IP", "Port"]
    widths1 = [Inches(1.0), Inches(1.8), Inches(1.8), Inches(1.0), Inches(1.1), Inches(0.6)]

    for col_idx, text in enumerate(headers1):
        cell = table1.cell(0, col_idx)
        cell.width = widths1[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        set_table_header_border(cell)

    data1 = [
        ["stu3_sys1", "Load Balancer & Generator", "ssh -p 2209 student@10.1.75.51", "10.1.75.51", "172.17.0.10", "4209"],
        ["stu3_sys2", "Backend 1 (Messaging App)", "ssh -p 2210 student@10.1.75.51", "10.1.75.51", "172.17.0.11", "4210"],
        ["stu3_sys3", "Backend 2 (Messaging App)", "ssh -p 2211 student@10.1.75.51", "10.1.75.51", "172.17.0.12", "4211"],
        ["stu3_sys4", "Backend 3 (Messaging App)", "ssh -p 2212 student@10.1.75.51", "10.1.75.51", "172.17.0.13", "4212"],
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table1.cell(row_idx, col_idx)
            cell.width = widths1[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            set_cell_border(cell)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(4)
    p_space.paragraph_format.space_after = Pt(4)

    add_para(
        "Communication Flow:\n"
        "- External Clients: Connect via http://10.1.75.51:4209/ and ws://10.1.75.51:4209/ws.\n"
        "- Internal Routing: The Load Balancer on stu3_sys1 (172.17.0.10) directly forwards HTTP and WebSocket traffic over the internal Docker network to 172.17.0.11:4210, 172.17.0.12:4211, and 172.17.0.13:4212."
    )

    # --- Section 3: Integration with Secure Messaging Application ---
    add_heading_1("3. Integration with the Secure Messaging Application")
    add_para(
        "The backend instances run the secure group messaging application developed in the previous assignment, "
        "built with Python, Flask, Flask-Sock, SQLite, and Cryptography. The key architectural and security layers include:"
    )
    add_para(
        "1. Encryption at Rest (Fernet / AES-128-CBC + HMAC-SHA256): All message plaintext is encrypted prior to database insertion. The SQLite database stores only Fernet ciphertext tokens.\n"
        "2. ECDSA Digital Signatures (P-256): Each user session creates an ECDSA keypair on curve SECP256R1. Messages are canonically formatted (username|text|timestamp) and signed with the sender's private key. Public keys are recorded in the database, allowing signature validation upon history retrieval.\n"
        "3. Tamper-Evident Hash Chain (SHA-256): Consecutive database entries are cryptographically linked: record_hash = SHA256(prev_hash || username || ciphertext || signature || timestamp). Modifying or deleting any database record breaks subsequent hashes, raising visual tamper warnings in the client UI.\n"
        "4. Real-Time WebSockets (/ws): Manages live bidirectional communication, presence detection, typing indicators, and message history distribution."
    )

    # --- Section 4: Go Load Balancer Architecture ---
    add_heading_1("4. Technical Architecture of the Go Load Balancer")
    add_para(
        "Traditional HTTP reverse proxies are designed for short-lived request-response cycles. WebSockets, however, "
        "require persistent, bidirectional TCP connections. The Go Load Balancer handles these requirements using the following design principles:"
    )
    add_heading_2("4.1 TCP Connection Hijacking and WebSocket Proxying")
    add_para(
        "When an incoming HTTP request contains Connection: Upgrade and Upgrade: websocket headers, the Load Balancer intercepts it using Go's http.Hijacker interface. "
        "It takes control of the underlying TCP socket, connects to the target backend via net.DialTimeout, and forwards the complete raw HTTP upgrade handshake (preserving Sec-WebSocket-Key and Sec-WebSocket-Version). "
        "Once established, two concurrent goroutines execute io.Copy in both directions, enabling full-duplex stream forwarding."
    )
    add_heading_2("4.2 Least Connections Routing Algorithm")
    add_para(
        "Because WebSocket sessions remain connected for extended periods, standard Round-Robin routing can lead to severe load imbalance. "
        "The Load Balancer implements Least Connections: it tracks active connections per backend atomically (sync/atomic). "
        "New connections are assigned to the backend holding the lowest active connection count. Counters are decremented when connections terminate."
    )
    add_heading_2("4.3 Active Health Checking")
    add_para(
        "A background goroutine polls the /health endpoint of each backend server every 2 seconds. "
        "If a backend fails to respond with HTTP 200 OK within 2 seconds, it is marked offline and excluded from routing until recovery."
    )

    # --- Section 5: Go Load Generator Architecture ---
    add_heading_1("5. Technical Architecture of the Go Load Generator")
    add_para(
        "The Load Generator (load_generator/main.go) is a standalone benchmarking tool written in Go without external dependencies. "
        "It implements the RFC-6455 WebSocket protocol directly:"
    )
    add_para(
        "- Concurrency: Spawns N concurrent goroutines (one per virtual user), each maintaining an open WebSocket connection to ws://172.17.0.10:4209/ws.\n"
        "- Protocol Framing: Constructs standard client-masked binary WebSocket text frames (opcode 0x81, random 4-byte mask key, XOR-masked payload).\n"
        "- Traffic Pattern: Sends join events followed by timestamped chat messages at 200 ms intervals.\n"
        "- Metrics Computation: Records microsecond-level message transit times, calculating throughput (msg/sec) and sorting latencies to determine minimum, average, median (p50), 95th percentile (p95), 99th percentile (p99), and maximum values."
    )

    # --- Section 6: Performance Benchmark Results ---
    add_heading_1("6. Empirical Performance Comparison and Results")
    add_para(
        "Benchmarks were conducted on stu3_sys1 with 50 concurrent virtual clients over a 20.56-second duration "
        "with a 200 ms message transmission interval per client."
    )

    # Table 2: Benchmark Comparison
    table2 = doc.add_table(rows=15, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False

    headers2 = ["Benchmark Metric", "Single Backend (Sys2 Only)", "Three Backends (Sys2+Sys3+Sys4)", "Variance / Analysis"]
    widths2 = [Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.8)]

    for col_idx, text in enumerate(headers2):
        cell = table2.cell(0, col_idx)
        cell.width = widths2[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9.5)
        set_table_header_border(cell)

    data2 = [
        ["Active Backend Nodes", "1 (Sys2:4210)", "3 (Sys2:4210, Sys3:4211, Sys4:4212)", "3x processing capacity"],
        ["Load Balancing Algorithm", "Passthrough", "Least Connections", "Balanced socket distribution"],
        ["Concurrent Clients", "50", "50", "High concurrent session load"],
        ["Test Duration", "20.56 s", "20.56 s", "Standardized testing window"],
        ["Total Messages Sent", "3,732", "4,644", "+24.44% messages processed"],
        ["Total Messages Received", "4,294", "7,647", "+78.08% broadcast delivery"],
        ["Effective Throughput", "181.48 msg/sec", "225.90 msg/sec", "+24.48% throughput increase"],
        ["Connection Errors / Drops", "29", "12", "-58.62% drop reduction"],
        ["Minimum Latency", "0.02 ms", "0.02 ms", "Equal baseline connection latency"],
        ["Average Latency", "0.26 ms", "0.32 ms", "Stable average response"],
        ["Median (p50) Latency", "0.12 ms", "0.13 ms", "Sub-millisecond median"],
        ["95th Percentile (p95)", "0.68 ms", "0.76 ms", "Sub-millisecond 95th percentile"],
        ["99th Percentile (p99)", "3.19 ms", "3.78 ms", "Consistent tail latency"],
        ["Maximum Latency", "48.35 ms", "37.73 ms", "-21.96% peak latency reduction"],
    ]

    for row_idx, row_data in enumerate(data2, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table2.cell(row_idx, col_idx)
            cell.width = widths2[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0 or "increase" in text or "reduction" in text:
                r.font.bold = True
            set_cell_border(cell)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(4)
    p_space2.paragraph_format.space_after = Pt(4)

    # --- Section 7: Results Discussion ---
    add_heading_1("7. Results and Discussion")
    add_para(
        "1. Throughput Scaling: Distributing WebSocket traffic across three backends increased throughput from 181.48 msg/sec to 225.90 msg/sec (+24.5%). "
        "Total broadcast delivery scaled from 4,294 to 7,647 messages (+78.1%), confirming that distributing ECDSA signing and Fernet encryption overhead relieves per-node CPU bottlenecks."
    )
    add_para(
        "2. Connection Reliability: Under high concurrency (50 simultaneous WebSocket handshakes), the single backend experienced 29 connection drops due to socket backlog limits. "
        "The 3-backend cluster reduced dropped connections to 12 (a 58.6% improvement in connection stability)."
    )
    add_para(
        "3. Peak Latency Reduction: Maximum latency dropped from 48.35 ms to 37.73 ms (-22.0%), demonstrating that the Least Connections algorithm effectively prevents thread starvation during heavy message bursts."
    )

    # --- Section 8: Load Balancer Source Code ---
    add_heading_1("8. Load Balancer Source Code (load_balancer/main.go)")
    
    code_text = open(r"c:\Users\lekka\Downloads\csd_proj\load_balancer\main.go", "r").read()
    
    p_code = doc.add_paragraph()
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(6)
    p_code.paragraph_format.line_spacing = 1.0
    r_code = p_code.add_run(code_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(8.5)

    # --- Section 9: Relevant Screenshots ---
    add_heading_1("9. Relevant Screenshots")

    img1 = r"C:/Users/lekka/.gemini/antigravity/brain/193ea9e5-1614-4c8f-9bd4-695d44fd37a4/.user_uploaded/media_1787581809116.png"
    img2 = r"C:/Users/lekka/.gemini/antigravity/brain/193ea9e5-1614-4c8f-9bd4-695d44fd37a4/.user_uploaded/media_1787581809309.png"
    img3 = r"C:/Users/lekka/.gemini/antigravity/brain/193ea9e5-1614-4c8f-9bd4-695d44fd37a4/.user_uploaded/media_1787581809122.png"

    # Screenshot 1
    add_heading_2("Figure 1: Go Load Balancer Execution and WebSocket Routing (stu3_sys1)")
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(6.2))
        p_cap1 = doc.add_paragraph("Figure 1: Load balancer runtime log showing registration of 172.17.0.11:4210, 172.17.0.12:4211, and 172.17.0.13:4212 and Least-Connections WebSocket stream forwarding.")
        p_cap1.paragraph_format.space_before = Pt(2)
        p_cap1.paragraph_format.space_after = Pt(10)
        p_cap1.runs[0].font.size = Pt(9.5)
        p_cap1.runs[0].font.italic = True

    # Screenshot 2
    add_heading_2("Figure 2: Single Backend Benchmark Results (Sys2 Only)")
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(6.2))
        p_cap2 = doc.add_paragraph("Figure 2: Load generator output with 50 concurrent WebSocket clients against Sys2 only (Throughput: 181.48 msg/sec, 29 drops, Max Latency: 48.35 ms).")
        p_cap2.paragraph_format.space_before = Pt(2)
        p_cap2.paragraph_format.space_after = Pt(10)
        p_cap2.runs[0].font.size = Pt(9.5)
        p_cap2.runs[0].font.italic = True

    # Screenshot 3
    add_heading_2("Figure 3: Three Backends Benchmark Results (Sys2 + Sys3 + Sys4)")
    if os.path.exists(img3):
        doc.add_picture(img3, width=Inches(6.2))
        p_cap3 = doc.add_paragraph("Figure 3: Load generator output with 50 concurrent WebSocket clients across all three backends (Throughput: 225.90 msg/sec, 12 drops, Max Latency: 37.73 ms).")
        p_cap3.paragraph_format.space_before = Pt(2)
        p_cap3.paragraph_format.space_after = Pt(10)
        p_cap3.runs[0].font.size = Pt(9.5)
        p_cap3.runs[0].font.italic = True

    output_path = r"c:\Users\lekka\Downloads\csd_proj\Report_12341330_Lekkala_Sashank.docx"
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_report()
